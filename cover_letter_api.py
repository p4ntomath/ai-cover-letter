import os
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, EmailStr
from typing import List, Optional
from docx import Document
from docx.shared import Pt
import uuid
from datetime import datetime
import uvicorn

# Load environment variables
load_dotenv()

app = FastAPI(
    title="Cover Letter Generator API",
    description="Generate professional cover letters in DOCX format",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Request Models
class CoverLetterRequest(BaseModel):
    file_name: Optional[str] = None
    your_name: str
    your_address: str
    your_email: EmailStr
    your_phone: str
    employer_name: str
    company_name: str
    company_address: str
    position_title: str
    body_paragraphs: List[str]

def generate_cover_letter_docx(data: CoverLetterRequest) -> str:
    """Generate cover letter DOCX file"""
    doc = Document()
    
    # Set font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Header - Your information
    doc.add_paragraph(data.your_name)
    doc.add_paragraph(data.your_address)
    doc.add_paragraph(f"Email: {data.your_email}")
    doc.add_paragraph(f"Phone: {data.your_phone}")
    doc.add_paragraph("")
    
    # Date
    today = datetime.today().strftime("%B %d, %Y")
    doc.add_paragraph(today)
    doc.add_paragraph("")
    
    # Employer information
    doc.add_paragraph(data.employer_name)
    doc.add_paragraph(data.company_name)
    doc.add_paragraph(data.company_address)
    doc.add_paragraph("")
    
    # Greeting
    doc.add_paragraph(f"Dear {data.employer_name},")
    doc.add_paragraph("")
    
    # Body paragraphs
    for paragraph in data.body_paragraphs:
        doc.add_paragraph(paragraph)
        doc.add_paragraph("")
    
    # Closing
    doc.add_paragraph("Thank you for considering my application.")
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(data.your_name)
    
    # Generate filename if not provided
    filename = data.file_name or f"cover_letter_{data.company_name}_{data.position_title}_{uuid.uuid4().hex[:8]}.docx"
    filename = filename.replace(" ", "_").replace("/", "_").replace("\\", "_")
    
    doc.save(filename)
    return filename

@app.get("/")
async def root():
    """API information"""
    return {
        "message": "Cover Letter Generator API",
        "version": "1.0.0",
        "endpoints": {
            "/generate": "POST - Generate cover letter DOCX file",
            "/health": "GET - Health check"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "message": "Cover letter generator is running"
    }

@app.post("/generate")
async def generate_cover_letter(data: CoverLetterRequest):
    """
    Generate a cover letter DOCX file from provided data
    
    Returns a downloadable DOCX file
    """
    try:
        # Generate the cover letter file
        filename = generate_cover_letter_docx(data)
        
        return FileResponse(
            path=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate cover letter: {str(e)}")

if __name__ == "__main__":
    print("Starting Cover Letter Generator API...")