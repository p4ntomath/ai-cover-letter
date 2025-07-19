import os
import pdfplumber
from docx import Document
import argparse
import json
from typing import Optional, Dict, List
from pathlib import Path

class TextExtractor:
    """
    A universal text extractor for PDF and DOCX files
    """
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_extension = Path(file_path).suffix.lower()
        self.filename = Path(file_path).stem
        
    def extract_from_pdf(self) -> Dict:
        """Extract text from PDF file"""
        try:
            text_content = ""
            page_texts = []
            metadata = {
                "total_pages": 0,
                "file_size": os.path.getsize(self.file_path),
                "file_type": "PDF"
            }
            
            with pdfplumber.open(self.file_path) as pdf:
                metadata["total_pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        page_texts.append({
                            "page_number": page_num,
                            "text": page_text.strip()
                        })
                        text_content += page_text + "\n\n"
                
                # Extract PDF metadata if available
                if pdf.metadata:
                    metadata.update({
                        "title": pdf.metadata.get('Title', ''),
                        "author": pdf.metadata.get('Author', ''),
                        "subject": pdf.metadata.get('Subject', ''),
                        "creator": pdf.metadata.get('Creator', ''),
                        "producer": pdf.metadata.get('Producer', ''),
                        "creation_date": str(pdf.metadata.get('CreationDate', '')),
                        "modification_date": str(pdf.metadata.get('ModDate', ''))
                    })
            
            return {
                "success": True,
                "filename": self.filename,
                "full_text": text_content.strip(),
                "pages": page_texts,
                "metadata": metadata,
                "word_count": len(text_content.split()),
                "character_count": len(text_content)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error extracting PDF: {str(e)}",
                "filename": self.filename
            }
    
    def extract_from_docx(self) -> Dict:
        """Extract text from DOCX file"""
        try:
            doc = Document(self.file_path)
            
            # Extract paragraph text
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    paragraphs.append(para_text)
                    full_text += para_text + "\n"
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                        full_text += cell_text + " "
                    table_data.append(row_data)
                table_texts.append(table_data)
            
            # Extract document properties
            core_props = doc.core_properties
            metadata = {
                "file_type": "DOCX",
                "file_size": os.path.getsize(self.file_path),
                "paragraph_count": len(paragraphs),
                "table_count": len(table_texts),
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or ""
            }
            
            return {
                "success": True,
                "filename": self.filename,
                "full_text": full_text.strip(),
                "paragraphs": paragraphs,
                "tables": table_texts,
                "metadata": metadata,
                "word_count": len(full_text.split()),
                "character_count": len(full_text)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error extracting DOCX: {str(e)}",
                "filename": self.filename
            }
    
    def extract_text(self) -> Dict:
        """Main method to extract text based on file type"""
        if not os.path.exists(self.file_path):
            return {
                "success": False,
                "error": f"File not found: {self.file_path}",
                "filename": self.filename
            }
        
        if self.file_extension == '.pdf':
            return self.extract_from_pdf()
        elif self.file_extension == '.docx':
            return self.extract_from_docx()
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {self.file_extension}. Supported types: .pdf, .docx",
                "filename": self.filename
            }
    
    def save_extracted_text(self, output_file: Optional[str] = None) -> str:
        """Extract text and save to file"""
        result = self.extract_text()
        
        if not result["success"]:
            print(f"❌ {result['error']}")
            return ""
        
        # Generate output filename if not provided
        if not output_file:
            output_file = f"{self.filename}_extracted.txt"
        
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(result["full_text"])
            
            print(f"✅ Text extracted and saved to: {output_file}")
            print(f"📄 Word count: {result['word_count']}")
            print(f"📊 Character count: {result['character_count']}")
            
            return output_file
            
        except Exception as e:
            print(f"❌ Error saving text: {e}")
            return ""
    
    def save_detailed_json(self, output_file: Optional[str] = None) -> str:
        """Extract all data and save as detailed JSON"""
        result = self.extract_text()
        
        if not output_file:
            output_file = f"{self.filename}_detailed.json"
        
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            
            if result["success"]:
                print(f"✅ Detailed extraction saved to: {output_file}")
            else:
                print(f"❌ Extraction failed: {result.get('error', 'Unknown error')}")
            
            return output_file
            
        except Exception as e:
            print(f"❌ Error saving JSON: {e}")
            return ""

def main():
    """Command line interface for text extraction"""
    parser = argparse.ArgumentParser(description='Extract text from PDF and DOCX files')
    parser.add_argument('file', help='Path to PDF or DOCX file')
    parser.add_argument('-o', '--output', help='Output text file path')
    parser.add_argument('-j', '--json', help='Save detailed JSON output')
    parser.add_argument('-t', '--text-only', action='store_true', help='Output only plain text')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.file):
        print(f"❌ File not found: {args.file}")
        return
    
    extractor = TextExtractor(args.file)
    
    print(f"🔍 Extracting text from: {args.file}")
    print("=" * 60)
    
    if args.text_only:
        # Just extract and print text
        result = extractor.extract_text()
        if result["success"]:
            print(result["full_text"])
        else:
            print(f"❌ Error: {result['error']}")
    else:
        # Extract and save to file
        if args.json:
            extractor.save_detailed_json(args.json)
        else:
            extractor.save_extracted_text(args.output)

if __name__ == "__main__":
    # If run directly with no args, show example usage
    import sys
    
    if len(sys.argv) == 1:
        print("📄 Text Extractor - PDF & DOCX")
        print("=" * 40)
        print("Extract text from PDF and DOCX files")
        print()
        print("Usage examples:")
        print("  python text_extractor.py document.pdf")
        print("  python text_extractor.py document.docx -o output.txt")
        print("  python text_extractor.py document.pdf -j detailed.json")
        print("  python text_extractor.py document.pdf --text-only")
        print()
        
        # Demo with existing files in the workspace
        print("🚀 Demo with files in current directory:")
        print("-" * 40)
        
        # Check for PDF files
        pdf_files = [f for f in os.listdir('.') if f.endswith('.pdf')]
        docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]
        
        if pdf_files:
            print(f"Found PDF file: {pdf_files[0]}")
            extractor = TextExtractor(pdf_files[0])
            result = extractor.extract_text()
            if result["success"]:
                print(f"✅ Successfully extracted {result['word_count']} words")
                print(f"📄 Preview: {result['full_text'][:200]}...")
            else:
                print(f"❌ Error: {result['error']}")
        
        if docx_files:
            print(f"\nFound DOCX file: {docx_files[0]}")
            extractor = TextExtractor(docx_files[0])
            result = extractor.extract_text()
            if result["success"]:
                print(f"✅ Successfully extracted {result['word_count']} words")
                print(f"📄 Preview: {result['full_text'][:200]}...")
            else:
                print(f"❌ Error: {result['error']}")
        
        if not pdf_files and not docx_files:
            print("No PDF or DOCX files found in current directory")
    else:
        main()