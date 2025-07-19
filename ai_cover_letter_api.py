import os
from dotenv import load_dotenv
from azure.ai.inference import ChatCompletionsClient
from azure.ai.inference.models import SystemMessage, UserMessage
from azure.core.credentials import AzureKeyCredential
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, EmailStr
from typing import List, Optional, Dict, Any
import json
import tempfile
import uvicorn
from pathlib import Path
import pdfplumber
from docx import Document
from docx.shared import Pt
import uuid
from datetime import datetime

# Load environment variables
load_dotenv()

app = FastAPI(
    title="AI-Powered Cover Letter Generator",
    description="Generate personalized cover letters using AI analysis of resumes and job descriptions",
    version="2.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

# GitHub AI Models setup
github_token = os.environ.get("GITHUB_TOKEN")

if github_token and github_token != "your_github_token_here":
    endpoint = "https://models.github.ai/inference"
    model = "gpt-4o"
    
    client = ChatCompletionsClient(
        endpoint=endpoint,
        credential=AzureKeyCredential(github_token),
    )
    print("✅ Using GitHub AI Models")
else:
    client = None
    print("❌ Warning: No valid GITHUB_TOKEN found")

# Response Models
class CoverLetterData(BaseModel):
    file_name: Optional[str] = None
    your_name: str
    your_address: str
    your_email: EmailStr
    your_phone: str
    employer_name: str
    company_name: str
    company_address: str
    position_title: str
    body_paragraphs: List[str]

class AIAnalysisResponse(BaseModel):
    success: bool
    extracted_data: Optional[CoverLetterData] = None
    error_message: Optional[str] = None
    ai_confidence: Optional[str] = None

class TextExtractor:
    """Text extraction utility"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        try:
            text_content = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n\n"
            return text_content.strip()
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        try:
            doc = Document(file_path)
            full_text = ""
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    full_text += para_text + "\n"
            return full_text.strip()
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")

class AIPromptEngineer:
    """AI Prompt Engineering for structured data extraction"""
    
    @staticmethod
    def create_system_prompt() -> str:
        return """You are a helpful assistant that extracts structured information from resumes and job descriptions for generating a personalized cover letter.

You must analyze the provided resume and job description, then extract the required information and generate professional cover letter body paragraphs.

IMPORTANT: Always return your response as valid JSON format with the exact structure requested. Do not include any markdown formatting or code blocks - just pure JSON."""

    @staticmethod
    def create_user_prompt(resume_text: str, job_description: str) -> str:
        return f"""Here is the resume:

{resume_text}

Here is the job description:

{job_description}

Please extract the following information and generate a personalized cover letter:

Required fields to extract/generate:
- file_name: Generate a suitable filename (e.g., "cover_letter_CompanyName_Position.docx")
- your_name: Extract from resume
- your_address: Extract from resume  
- your_email: Extract from resume
- your_phone: Extract from resume
- employer_name: Extract from job description (hiring manager name) or use "Hiring Manager"
- company_name: Extract from job description
- company_address: Extract from job description or use "Company Address"
- position_title: Extract from job description
- body_paragraphs: Generate 3-4 professional cover letter body paragraphs that:
  1. Show enthusiasm for the specific role and company
  2. Highlight relevant experience and skills from the resume that match the job requirements
  3. Demonstrate knowledge of the company/role from the job description
  4. Include specific examples of achievements that align with the job needs

IMPORTANT FORMATTING RULES:
- Do NOT use em dashes (—) anywhere in the text
- Use regular hyphens (-) for compound words and ranges
- Use commas, periods, and semicolons for punctuation
- Keep sentences clear and professional without special characters
- Use standard business letter formatting and language

Return the result in this exact JSON format:
{{
  "file_name": "cover_letter_example.docx",
  "your_name": "Full Name",
  "your_address": "Complete Address",
  "your_email": "email@example.com",
  "your_phone": "Phone Number",
  "employer_name": "Hiring Manager Name or 'Hiring Manager'",
  "company_name": "Company Name",
  "company_address": "Company Address",
  "position_title": "Job Title",
  "body_paragraphs": [
    "First paragraph expressing interest and mentioning how you learned about the position...",
    "Second paragraph highlighting relevant experience and skills...",
    "Third paragraph demonstrating company knowledge and cultural fit...",
    "Fourth paragraph with call to action and closing..."
  ]
}}"""

    @staticmethod
    def analyze_and_extract(resume_text: str, job_description: str) -> Dict[str, Any]:
        """Use AI to analyze resume and job description and extract structured data"""
        
        try:
            system_prompt = AIPromptEngineer.create_system_prompt()
            user_prompt = AIPromptEngineer.create_user_prompt(resume_text, job_description)
            
            # Use GitHub AI Models
            if not client:
                raise Exception("AI client not available - check GITHUB_TOKEN")
            
            print("Sending request to GitHub AI Models...")
            response = client.complete(
                messages=[
                    SystemMessage(system_prompt),
                    UserMessage(user_prompt),
                ],
                temperature=0.7,
                top_p=0.9,
                model=model
            )
            
            # Get the raw response
            ai_response = response.choices[0].message.content
            
            # Debug logging with more details
            print("AI Response received")
            print(f"Response type: {type(ai_response)}")
            print(f"Response length: {len(ai_response) if ai_response else 'None'}")
            print(f"First 200 chars: {ai_response[:200] if ai_response else 'None'}...")
            
            # Validate we got a response
            if not ai_response:
                return {
                    "success": False,
                    "error": "AI returned empty response",
                    "raw_response": str(ai_response)
                }
            
            # Clean up the response
            cleaned_response = ai_response.strip()
            
            # Remove markdown code blocks if present
            if cleaned_response.startswith('```json'):
                cleaned_response = cleaned_response.replace('```json', '').replace('```', '').strip()
            elif cleaned_response.startswith('```'):
                cleaned_response = cleaned_response.replace('```', '').strip()
            
            # Remove any surrounding quotes if the entire response is quoted
            if cleaned_response.startswith('"') and cleaned_response.endswith('"') and cleaned_response.count('"') == 2:
                cleaned_response = cleaned_response[1:-1]
            
            print(f"Cleaned response length: {len(cleaned_response)}")
            print(f"Cleaned first 200 chars: {cleaned_response[:200]}...")
            
            # Try to parse as JSON
            try:
                extracted_data = json.loads(cleaned_response)
                print("JSON parsing successful")
                print(f"Extracted data type: {type(extracted_data)}")
                
                # Validate that we got a dictionary
                if not isinstance(extracted_data, dict):
                    return {
                        "success": False,
                        "error": f"AI returned {type(extracted_data).__name__} instead of expected dictionary",
                        "raw_response": ai_response,
                        "cleaned_response": cleaned_response
                    }
                
                # Validate required fields are present
                required_fields = ['your_name', 'your_email', 'your_phone', 'company_name', 'position_title', 'body_paragraphs']
                missing_fields = [field for field in required_fields if field not in extracted_data]
                
                if missing_fields:
                    return {
                        "success": False,
                        "error": f"Missing required fields: {missing_fields}",
                        "raw_response": ai_response,
                        "extracted_data": extracted_data
                    }
                
                print("All required fields present")
                return {
                    "success": True,
                    "data": extracted_data,
                    "confidence": "high"
                }
                
            except json.JSONDecodeError as e:
                print(f"JSON parsing failed: {str(e)}")
                
                # Try to extract JSON from within the response
                try:
                    # Look for JSON-like content between curly braces
                    import re
                    json_match = re.search(r'\{.*\}', cleaned_response, re.DOTALL)
                    if json_match:
                        json_content = json_match.group(0)
                        print(f"Found JSON content: {json_content[:200]}...")
                        extracted_data = json.loads(json_content)
                        
                        if isinstance(extracted_data, dict):
                            print("JSON extraction successful from regex")
                            return {
                                "success": True,
                                "data": extracted_data,
                                "confidence": "medium"
                            }
                except:
                    pass
                
                return {
                    "success": False,
                    "error": f"Failed to parse AI response as JSON: {str(e)}",
                    "raw_response": ai_response,
                    "cleaned_response": cleaned_response,
                    "json_error": str(e)
                }
                
        except Exception as e:
            print(f"AI analysis exception: {str(e)}")
            return {
                "success": False,
                "error": f"AI analysis failed: {str(e)}"
            }

def generate_cover_letter_docx(data: CoverLetterData) -> str:
    """Generate cover letter DOCX file"""
    doc = Document()
    
    # Set font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Header - Your information
    doc.add_paragraph(data.your_name)
    doc.add_paragraph(data.your_address)
    doc.add_paragraph(f"Email: {data.your_email}")
    doc.add_paragraph(f"Phone: {data.your_phone}")
    doc.add_paragraph("")
    
    # Date
    today = datetime.today().strftime("%B %d, %Y")
    doc.add_paragraph(today)
    doc.add_paragraph("")
    
    # Employer information
    doc.add_paragraph(data.employer_name)
    doc.add_paragraph(data.company_name)
    doc.add_paragraph(data.company_address)
    doc.add_paragraph("")
    
    # Greeting
    doc.add_paragraph(f"Dear {data.employer_name},")
    doc.add_paragraph("")
    
    # Body paragraphs
    for paragraph in data.body_paragraphs:
        doc.add_paragraph(paragraph)
        doc.add_paragraph("")
    
    # Closing
    doc.add_paragraph("Thank you for considering my application.")
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(data.your_name)
    
    # Generate filename if not provided
    filename = data.file_name or f"cover_letter_{data.company_name}_{data.position_title}_{uuid.uuid4().hex[:8]}.docx"
    filename = filename.replace(" ", "_").replace("/", "_").replace("\\", "_")
    
    doc.save(filename)
    return filename

@app.get("/")
async def root():
    """API information"""
    return {
        "message": "AI-Powered Cover Letter Generator",
        "version": "2.0.0",
        "features": [
            "AI analysis of resumes and job descriptions",
            "Automatic data extraction",
            "Personalized cover letter generation",
            "DOCX file output"
        ],
        "endpoints": {
            "/generate-ai-cover-letter": "POST - Upload resume + job description for AI analysis",
            "/analyze-documents": "POST - Analyze documents and return extracted data (no file generation)",
            "/health": "GET - Health check"
        }
    }

@app.get("/health")
async def health_check():
    """Health check"""
    if client:
        ai_status = "✅ GitHub AI Models available"
    else:
        ai_status = "❌ No AI service available (check GITHUB_TOKEN)"
    
    return {
        "status": "healthy",
        "ai_service": ai_status,
        "message": "AI-powered cover letter generator is running"
    }

@app.post("/analyze-documents", response_model=AIAnalysisResponse)
async def analyze_documents(
    resume: UploadFile = File(...),
    job_description: UploadFile = File(None),
    job_description_text: str = Form(None)
):
    """
    Analyze resume and job description using AI to extract structured data
    
    - **resume**: Upload resume file (PDF or DOCX)
    - **job_description**: Upload job description file (PDF or DOCX) OR
    - **job_description_text**: Provide job description as text
    """
    
    if not client:
        raise HTTPException(status_code=503, detail="AI service unavailable - missing GITHUB_TOKEN")
    
    try:
        # Handle resume text extraction
        resume_filename = resume.filename or "resume.txt"
        resume_ext = Path(resume_filename).suffix.lower()
        
        # If it's a text file (from HTML blob), read directly
        if resume_ext in ['.txt', ''] or resume_filename == 'resume.txt':
            resume_content = await resume.read()
            resume_text = resume_content.decode('utf-8')
        else:
            # Handle PDF/DOCX files
            if resume_ext not in ['.pdf', '.docx']:
                raise HTTPException(status_code=400, detail="Resume must be PDF, DOCX, or text")
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=resume_ext) as temp_resume:
                resume_content = await resume.read()
                temp_resume.write(resume_content)
                temp_resume_path = temp_resume.name
            
            if resume_ext == '.pdf':
                resume_text = TextExtractor.extract_from_pdf(temp_resume_path)
            else:
                resume_text = TextExtractor.extract_from_docx(temp_resume_path)
            
            os.unlink(temp_resume_path)
        
        # Get job description text
        if job_description_text:
            job_desc_text = job_description_text
        elif job_description:
            job_ext = Path(job_description.filename).suffix.lower()
            if job_ext not in ['.pdf', '.docx']:
                raise HTTPException(status_code=400, detail="Job description must be PDF or DOCX")
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=job_ext) as temp_job:
                job_content = await job_description.read()
                temp_job.write(job_content)
                temp_job_path = temp_job.name
            
            if job_ext == '.pdf':
                job_desc_text = TextExtractor.extract_from_pdf(temp_job_path)
            else:
                job_desc_text = TextExtractor.extract_from_docx(temp_job_path)
            
            os.unlink(temp_job_path)
        else:
            raise HTTPException(status_code=400, detail="Must provide either job_description file or job_description_text")
        
        # Validate we have content
        if not resume_text.strip():
            raise HTTPException(status_code=400, detail="Resume text is empty")
        if not job_desc_text.strip():
            raise HTTPException(status_code=400, detail="Job description text is empty")
        
        # AI Analysis
        result = AIPromptEngineer.analyze_and_extract(resume_text, job_desc_text)
        
        if result["success"]:
            try:
                cover_letter_data = CoverLetterData(**result["data"])
                return AIAnalysisResponse(
                    success=True,
                    extracted_data=cover_letter_data,
                    ai_confidence=result.get("confidence", "unknown")
                )
            except Exception as e:
                return AIAnalysisResponse(
                    success=False,
                    error_message=f"Data validation error: {str(e)}"
                )
        else:
            return AIAnalysisResponse(
                success=False,
                error_message=result.get("error", "Unknown AI analysis error")
            )
            
    except HTTPException:
        raise
    except Exception as e:
        return AIAnalysisResponse(
            success=False,
            error_message=f"Analysis failed: {str(e)}"
        )

@app.post("/generate-ai-cover-letter")
async def generate_ai_cover_letter(
    resume: UploadFile = File(...),
    job_description: UploadFile = File(None),
    job_description_text: str = Form(None)
):
    """
    Generate a complete cover letter using AI analysis of resume and job description
    
    - **resume**: Upload resume file (PDF or DOCX)  
    - **job_description**: Upload job description file (PDF or DOCX) OR
    - **job_description_text**: Provide job description as text
    
    Returns a downloadable DOCX cover letter file
    """
    
    # First analyze the documents
    analysis_result = await analyze_documents(resume, job_description, job_description_text)
    
    if not analysis_result.success:
        raise HTTPException(status_code=400, detail=analysis_result.error_message)
    
    try:
        # Generate the cover letter file
        filename = generate_cover_letter_docx(analysis_result.extracted_data)
        
        return FileResponse(
            path=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate cover letter: {str(e)}")

if __name__ == "__main__":
    print("Starting AI-Powered Cover Letter Generator...")
    print("API will be available at: http://localhost:8003")
    print("API documentation at: http://localhost:8003/docs")
    print(f"AI Service: {'✅ Available' if client else '❌ Unavailable (missing GITHUB_TOKEN)'}")
    uvicorn.run(app, host="0.0.0.0", port=8003)