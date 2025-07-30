from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, Dict, List
import os
import tempfile
import pdfplumber
from docx import Document
from pathlib import Path

# Initialize FastAPI app
app = FastAPI(
    title="Text Extractor API",
    description="API to extract text from PDF and DOCX files",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

# Response models
class TextExtractionResponse(BaseModel):
    success: bool
    filename: str
    full_text: Optional[str] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    file_type: Optional[str] = None
    error_message: Optional[str] = None

class DetailedExtractionResponse(BaseModel):
    success: bool
    filename: str
    full_text: Optional[str] = None
    pages: Optional[List[Dict]] = None  # For PDF pages
    paragraphs: Optional[List[str]] = None  # For DOCX paragraphs
    tables: Optional[List] = None  # For DOCX tables
    metadata: Optional[Dict] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    error_message: Optional[str] = None

class TextExtractorAPI:
    """Text extraction logic for the API"""
    
    @staticmethod
    def extract_from_pdf(file_path: str, filename: str) -> Dict:
        """Extract text from PDF file"""
        try:
            text_content = ""
            page_texts = []
            metadata = {
                "file_type": "PDF",
                "total_pages": 0,
                "file_size": os.path.getsize(file_path)
            }
            
            with pdfplumber.open(file_path) as pdf:
                metadata["total_pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        page_texts.append({
                            "page_number": page_num,
                            "text": page_text.strip()
                        })
                        text_content += page_text + "\n\n"
                
                # Extract PDF metadata if available
                if pdf.metadata:
                    metadata.update({
                        "title": pdf.metadata.get('Title', ''),
                        "author": pdf.metadata.get('Author', ''),
                        "subject": pdf.metadata.get('Subject', ''),
                        "creator": pdf.metadata.get('Creator', ''),
                        "producer": pdf.metadata.get('Producer', '')
                    })
            
            return {
                "success": True,
                "filename": filename,
                "full_text": text_content.strip(),
                "pages": page_texts,
                "metadata": metadata,
                "word_count": len(text_content.split()),
                "character_count": len(text_content),
                "file_type": "PDF"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error_message": f"Error extracting PDF: {str(e)}",
                "filename": filename
            }
    
    @staticmethod
    def extract_from_docx(file_path: str, filename: str) -> Dict:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract paragraph text
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    paragraphs.append(para_text)
                    full_text += para_text + "\n"
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                        full_text += cell_text + " "
                    table_data.append(row_data)
                table_texts.append(table_data)
            
            # Extract document properties
            core_props = doc.core_properties
            metadata = {
                "file_type": "DOCX",
                "file_size": os.path.getsize(file_path),
                "paragraph_count": len(paragraphs),
                "table_count": len(table_texts),
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or ""
            }
            
            return {
                "success": True,
                "filename": filename,
                "full_text": full_text.strip(),
                "paragraphs": paragraphs,
                "tables": table_texts,
                "metadata": metadata,
                "word_count": len(full_text.split()),
                "character_count": len(full_text),
                "file_type": "DOCX"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error_message": f"Error extracting DOCX: {str(e)}",
                "filename": filename
            }

@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Text Extractor API",
        "version": "1.0.0",
        "supported_formats": ["PDF", "DOCX"],
        "endpoints": {
            "/extract": "POST - Extract text from uploaded file (simple response)",
            "/extract-detailed": "POST - Extract text with detailed structure",
            "/health": "GET - Health check endpoint",
            "/docs": "GET - API documentation"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "message": "Text extractor API is running"}

@app.post("/extract", response_model=TextExtractionResponse)
async def extract_text(file: UploadFile = File(...)):
    """
    Extract text from uploaded PDF or DOCX file (simple response)
    
    - **file**: Upload a PDF or DOCX file
    
    Returns extracted text with basic information
    """
    
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in ['.pdf', '.docx']:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type: {file_extension}. Supported types: .pdf, .docx"
        )
    
    # Create temporary file
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        # Extract text based on file type
        if file_extension == '.pdf':
            result = TextExtractorAPI.extract_from_pdf(temp_file_path, file.filename)
        else:  # .docx
            result = TextExtractorAPI.extract_from_docx(temp_file_path, file.filename)
        
        # Clean up temporary file
        os.unlink(temp_file_path)
        
        # Return simple response
        if result["success"]:
            return TextExtractionResponse(
                success=True,
                filename=result["filename"],
                full_text=result["full_text"],
                word_count=result["word_count"],
                character_count=result["character_count"],
                file_type=result["file_type"]
            )
        else:
            return TextExtractionResponse(
                success=False,
                filename=file.filename,
                error_message=result["error_message"]
            )
            
    except Exception as e:
        # Clean up temporary file if it exists
        if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.post("/extract-detailed", response_model=DetailedExtractionResponse)
async def extract_text_detailed(file: UploadFile = File(...)):
    """
    Extract text from uploaded PDF or DOCX file (detailed response)
    
    - **file**: Upload a PDF or DOCX file
    
    Returns extracted text with detailed structure (pages, paragraphs, tables, metadata)
    """
    
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in ['.pdf', '.docx']:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type: {file_extension}. Supported types: .pdf, .docx"
        )
    
    # Create temporary file
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        # Extract text based on file type
        if file_extension == '.pdf':
            result = TextExtractorAPI.extract_from_pdf(temp_file_path, file.filename)
        else:  # .docx
            result = TextExtractorAPI.extract_from_docx(temp_file_path, file.filename)
        
        # Clean up temporary file
        os.unlink(temp_file_path)
        
        # Return detailed response
        if result["success"]:
            return DetailedExtractionResponse(
                success=True,
                filename=result["filename"],
                full_text=result["full_text"],
                pages=result.get("pages"),
                paragraphs=result.get("paragraphs"),
                tables=result.get("tables"),
                metadata=result["metadata"],
                word_count=result["word_count"],
                character_count=result["character_count"]
            )
        else:
            return DetailedExtractionResponse(
                success=False,
                filename=file.filename,
                error_message=result["error_message"]
            )
            
    except Exception as e:
        # Clean up temporary file if it exists
        if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.post("/extract-text-only")
async def extract_text_only(file: UploadFile = File(...)):
    """
    Extract only the plain text from uploaded file (lightweight response)
    
    - **file**: Upload a PDF or DOCX file
    
    Returns only the extracted text as a string
    """
    
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in ['.pdf', '.docx']:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type: {file_extension}. Supported types: .pdf, .docx"
        )
    
    # Create temporary file
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        # Extract text based on file type
        if file_extension == '.pdf':
            result = TextExtractorAPI.extract_from_pdf(temp_file_path, file.filename)
        else:  # .docx
            result = TextExtractorAPI.extract_from_docx(temp_file_path, file.filename)
        
        # Clean up temporary file
        os.unlink(temp_file_path)
        
        # Return just the text
        if result["success"]:
            return {"text": result["full_text"]}
        else:
            raise HTTPException(status_code=400, detail=result["error_message"])
            
    except HTTPException:
        raise
    except Exception as e:
        # Clean up temporary file if it exists
        if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

if __name__ == "__main__":
    print("Starting Text Extractor API...")
    print("Supported file types: PDF, DOCX")